"""
Script to generate a highly professional Microsoft Word (.docx) Technical Document
for PreHub (Dokumen Pendukung Teknis) with embedded Playwright screenshots.
"""

import os
import sys
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, fill_hex):
    """Set background color of a table cell."""
    tcPr = cell._element.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    """Set inner margins (padding) of a table cell in dxa (1 pt = 20 dxa)."""
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = parse_xml(
        f'<w:tcMar {nsdecls("w")}>'
        f'<w:top w:w="{top}" w:type="dxa"/>'
        f'<w:bottom w:w="{bottom}" w:type="dxa"/>'
        f'<w:left w:w="{left}" w:type="dxa"/>'
        f'<w:right w:w="{right}" w:type="dxa"/>'
        f'</w:tcMar>'
    )
    tcPr.append(tcMar)

def set_cell_borders(cell, top=None, bottom=None, left=None, right=None):
    """Set borders on a cell."""
    tcPr = cell._element.get_or_add_tcPr()
    for child in list(tcPr):
        if child.tag.endswith('tcBorders'):
            tcPr.remove(child)
            
    tcBorders = OxmlElement('w:tcBorders')
    for side_name, side_cfg in [('top', top), ('left', left), ('bottom', bottom), ('right', right)]:
        edge = OxmlElement(f'w:{side_name}')
        if side_cfg:
            edge.set(qn('w:val'), side_cfg.get('val', 'single'))
            edge.set(qn('w:sz'), str(side_cfg.get('sz', '4')))
            edge.set(qn('w:space'), '0')
            edge.set(qn('w:color'), side_cfg.get('color', 'auto'))
        else:
            edge.set(qn('w:val'), 'none')
        tcBorders.append(edge)
    tcPr.append(tcBorders)

def add_header_footer(doc):
    """Add modern headers and footers with page numbering."""
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)

    # Header
    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    hrun = hp.add_run("PREHUB — DOKUMEN PENDUKUNG TEKNIS (TECHNICAL DOCUMENT)")
    hrun.font.name = "Calibri"
    hrun.font.size = Pt(8.5)
    hrun.font.color.rgb = RGBColor(148, 163, 184) # Slate 400

    # Footer
    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    frun1 = fp.add_run("PIDI — Digdaya x Hackathon 2026 | Bank Indonesia & OJK | Sistem PreHub v1.2.0")
    frun1.font.name = "Calibri"
    frun1.font.size = Pt(8.5)
    frun1.font.color.rgb = RGBColor(148, 163, 184)

def create_styled_heading1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = RGBColor(30, 58, 138) # Deep Navy (#1E3A8A)
    return p

def create_styled_heading2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.color.rgb = RGBColor(2, 132, 199) # Slate Cyan (#0284C7)
    return p

def create_styled_heading3(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(10.5)
    run.font.bold = True
    run.font.color.rgb = RGBColor(51, 65, 85) # Slate 700
    return p

def add_body_p(doc, text, bold_prefix=None, space_after=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.15
    if bold_prefix:
        brun = p.add_run(bold_prefix)
        brun.font.name = "Calibri"
        brun.font.size = Pt(10.5)
        brun.font.bold = True
        brun.font.color.rgb = RGBColor(30, 41, 59)
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(10.5)
    run.font.color.rgb = RGBColor(51, 65, 85)
    return p

def add_bullet_p(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.15
    if bold_prefix:
        brun = p.add_run(bold_prefix)
        brun.font.name = "Calibri"
        brun.font.size = Pt(10)
        brun.font.bold = True
        brun.font.color.rgb = RGBColor(30, 41, 59)
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(51, 65, 85)
    return p

def add_code_block(doc, code_text):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    
    cell = table.cell(0, 0)
    cell.width = Inches(6.5)
    set_cell_background(cell, "F8FAFC")
    set_cell_margins(cell, top=120, bottom=120, left=160, right=160)
    set_cell_borders(
        cell,
        top={"val": "single", "sz": "4", "color": "CBD5E1"},
        bottom={"val": "single", "sz": "4", "color": "CBD5E1"},
        left={"val": "single", "sz": "18", "color": "0284C7"}, # Blue accent border
        right={"val": "single", "sz": "4", "color": "CBD5E1"}
    )
    
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.05
    run = p.add_run(code_text)
    run.font.name = "Consolas"
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor(30, 41, 59)
    
    # Empty spacing paragraph
    sp = doc.add_paragraph()
    sp.paragraph_format.space_before = Pt(0)
    sp.paragraph_format.space_after = Pt(4)

def add_callout_box(doc, title, text, box_type="info"):
    colors = {
        "info": {"bg": "F0F9FF", "border": "0284C7", "title_color": RGBColor(2, 132, 199)},
        "success": {"bg": "F0FDF4", "border": "16A34A", "title_color": RGBColor(22, 163, 74)},
        "warning": {"bg": "FFFBEB", "border": "D97706", "title_color": RGBColor(217, 119, 6)},
    }
    cfg = colors.get(box_type, colors["info"])
    
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    
    cell = table.cell(0, 0)
    cell.width = Inches(6.5)
    set_cell_background(cell, cfg["bg"])
    set_cell_margins(cell, top=120, bottom=120, left=160, right=160)
    set_cell_borders(
        cell,
        top={"val": "single", "sz": "4", "color": "E2E8F0"},
        bottom={"val": "single", "sz": "4", "color": "E2E8F0"},
        left={"val": "single", "sz": "24", "color": cfg["border"]},
        right={"val": "single", "sz": "4", "color": "E2E8F0"}
    )
    
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.1
    
    trun = p.add_run(f"📌 {title}\n")
    trun.font.name = "Arial"
    trun.font.size = Pt(10)
    trun.font.bold = True
    trun.font.color.rgb = cfg["title_color"]
    
    mrun = p.add_run(text)
    mrun.font.name = "Calibri"
    mrun.font.size = Pt(9.5)
    mrun.font.color.rgb = RGBColor(51, 65, 85)
    
    sp = doc.add_paragraph()
    sp.paragraph_format.space_before = Pt(0)
    sp.paragraph_format.space_after = Pt(4)

def add_styled_table(doc, headers, rows_data, col_widths=None):
    table = doc.add_table(rows=len(rows_data) + 1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    
    # Header Row
    hdr_cells = table.rows[0].cells
    for i, title in enumerate(headers):
        hdr_cells[i].text = title
        set_cell_background(hdr_cells[i], "1E3A8A")
        set_cell_margins(hdr_cells[i], top=100, bottom=100, left=120, right=120)
        set_cell_borders(
            hdr_cells[i],
            top={"val": "single", "sz": "4", "color": "1E3A8A"},
            bottom={"val": "single", "sz": "12", "color": "0284C7"},
            left={"val": "single", "sz": "4", "color": "3B82F6"},
            right={"val": "single", "sz": "4", "color": "3B82F6"}
        )
        p = hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for run in p.runs:
            run.font.name = "Arial"
            run.font.size = Pt(9.5)
            run.font.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
            
    # Data Rows
    for r_idx, row_values in enumerate(rows_data):
        row_cells = table.rows[r_idx + 1].cells
        bg_color = "F8FAFC" if r_idx % 2 == 1 else "FFFFFF"
        for c_idx, val in enumerate(row_values):
            row_cells[c_idx].text = val
            set_cell_background(row_cells[c_idx], bg_color)
            set_cell_margins(row_cells[c_idx], top=80, bottom=80, left=120, right=120)
            set_cell_borders(
                row_cells[c_idx],
                top={"val": "single", "sz": "4", "color": "E2E8F0"},
                bottom={"val": "single", "sz": "4", "color": "E2E8F0"},
                left={"val": "single", "sz": "4", "color": "E2E8F0"},
                right={"val": "single", "sz": "4", "color": "E2E8F0"}
            )
            p = row_cells[c_idx].paragraphs[0]
            for run in p.runs:
                run.font.name = "Calibri"
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(51, 65, 85)
                
    # Column Widths
    if col_widths:
        for row in table.rows:
            for idx, width in enumerate(col_widths):
                row.cells[idx].width = width

    sp = doc.add_paragraph()
    sp.paragraph_format.space_before = Pt(0)
    sp.paragraph_format.space_after = Pt(4)

def add_screenshot_figure(doc, image_path, caption_title, description_text):
    if not os.path.exists(image_path):
        add_callout_box(doc, f"Image Not Found: {os.path.basename(image_path)}", "Gambar tidak ditemukan pada direktori lokal.", "warning")
        return
        
    p_img = doc.add_paragraph()
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_img.paragraph_format.space_before = Pt(8)
    p_img.paragraph_format.space_after = Pt(4)
    p_img.paragraph_format.keep_with_next = True
    
    run_img = p_img.add_run()
    run_img.add_picture(image_path, width=Inches(6.2))
    
    p_cap = doc.add_paragraph()
    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_cap.paragraph_format.space_before = Pt(2)
    p_cap.paragraph_format.space_after = Pt(4)
    p_cap.paragraph_format.keep_with_next = True
    
    crun1 = p_cap.add_run(f"Gambar: {caption_title} — ")
    crun1.font.name = "Arial"
    crun1.font.size = Pt(9)
    crun1.font.bold = True
    crun1.font.color.rgb = RGBColor(30, 58, 138)
    
    crun2 = p_cap.add_run(description_text)
    crun2.font.name = "Calibri"
    crun2.font.size = Pt(9)
    crun2.font.italic = True
    crun2.font.color.rgb = RGBColor(100, 116, 139)
    
    sp = doc.add_paragraph()
    sp.paragraph_format.space_before = Pt(0)
    sp.paragraph_format.space_after = Pt(4)


def build_technical_document(output_path, project_root):
    doc = Document()
    add_header_footer(doc)
    
    # -------------------------------------------------------------
    # COVER / HEADER SECTION
    # -------------------------------------------------------------
    title_p = doc.add_paragraph()
    title_p.paragraph_format.space_before = Pt(10)
    title_p.paragraph_format.space_after = Pt(4)
    run_t = title_p.add_run("DOKUMEN PENDUKUNG TEKNIS (TECHNICAL DOCUMENT)")
    run_t.font.name = "Arial"
    run_t.font.size = Pt(18)
    run_t.font.bold = True
    run_t.font.color.rgb = RGBColor(30, 58, 138) # Navy

    sub_p = doc.add_paragraph()
    sub_p.paragraph_format.space_before = Pt(0)
    sub_p.paragraph_format.space_after = Pt(12)
    run_sub = sub_p.add_run("SISTEM PREHUB: EARLY WARNING & MITIGATION DECISION SUPPORT SYSTEM UNTUK DISTRIBUSI PANGAN BERBASIS MULTI-AGENT SWARM DAN DATA MULTISUMBER")
    run_sub.font.name = "Arial"
    run_sub.font.size = Pt(11)
    run_sub.font.bold = True
    run_sub.font.color.rgb = RGBColor(2, 132, 199) # Cyan

    # Metadata Table
    meta_headers = ["Properti Dokumen", "Keterangan Spesifikasi"]
    meta_rows = [
        ["Nama Sistem / Produk", "PreHub (Predictive Logistics Hub & Early Warning Decision Support System)"],
        ["Versi Rilis / Milestone", "MVP v1.2.0-PROD (Phase 34 Multi-Modal Sumatra Release)"],
        ["Kategori Inovasi", "Sistem Pendukung Keputusan (Decision Support System - DSS) / AI-Driven Geo-Logistics"],
        ["Wilayah Operasional", "Koridor Strategis Pulau Sumatera (Darat, Laut Selat Malaka/Sunda, Udara KNO-CGK)"],
        ["Target Pengguna Utama", "Badan Pangan Nasional (BAPANAS), Kementerian Perhubungan (Kemenhub), Perum BULOG, Dishub/POLRI, Dispatcher Logistik"],
        ["Institusi Penyelenggara", "Inisiasi Pusat Inovasi Digital Indonesia (PIDI) — Digdaya x Hackathon 2026 (Bank Indonesia & OJK)"],
        ["Tanggal Dokumen", "17 Agustus 2026"]
    ]
    add_styled_table(doc, meta_headers, meta_rows, col_widths=[Inches(2.2), Inches(4.3)])

    add_callout_box(
        doc,
        "Dokumen Resmi Pendukung Teknis (Technical Installation & User Guide)",
        "Dokumen ini memuat panduan instalasi teknis, arsitektur struktural sistem, deskripsi fungsional modul, panduan operasional pengguna (SOP), formulasi matematika, pembuktian visual (tangkapan layar otomatis Playwright), dan rekomendasi arsitektur hosting cloud gratis ($0/bulan).",
        "info"
    )

    # -------------------------------------------------------------
    # BAB 1: RINGKASAN EKSEKUTIF & LATAR BELAKANG
    # -------------------------------------------------------------
    create_styled_heading1(doc, "BAB 1: RINGKASAN EKSEKUTIF & LATAR BELAKANG SISTEM")
    
    create_styled_heading2(doc, "1.1 Problem Statement")
    add_body_p(doc, "Distribusi komoditas pangan pokok (beras, cabai merah, aneka bawang, minyak goreng, dan daging) di Indonesia khususnya koridor Pulau Sumatera menghadapi tiga kerentanan struktural utama:")
    add_bullet_p(doc, " Banjir rob pesisir (Belawan), longsor lereng bukit (Jalan Lintas Sumatera), dan gelombang tinggi laut yang melumpuhkan rute pelayaran.", "1. Cuaca Ekstrem & Hidrometeorologi:")
    add_bullet_p(doc, " Antrean bongkar muat pelabuhan utama, kemacetan simpang arteri non-tol, jembatan rusak, atau insiden kecelakaan kendaraan berat.", "2. Kemacetan Kritis & Bottleneck:")
    add_bullet_p(doc, " Keterlambatan pengiriman komoditas pangan basah (perishable food) menyebabkan penyusutan bobot, kebusukan muatan (spoilage), dan lonjakan inflasi pangan lokal (volatile food inflation).", "3. Volatilitas Harga & Disparitas Spasial:")

    create_styled_heading2(doc, "1.2 Solusi PreHub")
    add_body_p(doc, "PreHub hadir sebagai platform intelijen logistik pangan terpadu (Unified Food Logistics Command Center) yang mengintegrasikan:")
    add_bullet_p(doc, " Mengintegrasikan data gempa/cuaca BMKG, proyeksi presipitasi Open-Meteo 48 jam, telemetri TomTom Traffic, dan intelijen berita regional Google News RSS NLP.", "• Multi-Source Data Grounding:")
    add_bullet_p(doc, " Kolaborasi 6 agen AI berbasis LangGraph (Data Collection, OSINT Hazard, Weather/Traffic Forecast, Route Optimization, Economic Intelligence, dan Decision Copilot DeepSeek R1).", "• Multi-Agent Swarm Architecture:")
    add_bullet_p(doc, " Pemodelan perutean komprehensif darat (arteri & tol), laut (Tol Laut Selat Malaka & Selat Sunda), serta udara (kargo KNO-CGK).", "• Multi-Modal Network Routing:")
    add_bullet_p(doc, " Rekomendasi mitigasi berbasis bukti (Evidence Chain) dengan 3 opsi aksi terukur: Continue (lanjutkan), Reroute (alihkan rute bypass), atau Hold/Delay (tahan di buffer depot terdekat).", "• Actionable Decision Support:")

    # -------------------------------------------------------------
    # BAB 2: PERSYARATAN SISTEM
    # -------------------------------------------------------------
    create_styled_heading1(doc, "BAB 2: PERSYARATAN SISTEM (SYSTEM REQUIREMENTS)")
    
    create_styled_heading2(doc, "2.1 Hardware Requirements")
    hw_headers = ["Komponen", "Server Minimum (Demo/Staging)", "Server Rekomendasi (Production)", "Workstation Klien / Dispatcher"]
    hw_rows = [
        ["Processor (CPU)", "2 Cores @ 2.0 GHz (x86_64 / ARM64)", "8-16 Cores @ 3.2 GHz (AMD EPYC / Xeon)", "4 Cores @ 2.0 GHz"],
        ["Memory (RAM)", "4 GB DDR4", "16 - 32 GB DDR4/DDR5", "8 GB DDR4"],
        ["Storage (Disk)", "10 GB SSD NVMe", "50 GB SSD NVMe (RAID 1)", "5 GB Ruang Kosong"],
        ["Graphics (GPU)", "Opsional (CPU Dijkstra Mode)", "NVIDIA T4 / RTX 4000 (cuOpt CUDA)", "GPU Terintegrasi (WebGL 2.0)"],
        ["Jaringan", "10 Mbps Dedicated", "100 Mbps Dedicated Full-Duplex", "5 Mbps Internet Stabil"]
    ]
    add_styled_table(doc, hw_headers, hw_rows, col_widths=[Inches(1.5), Inches(1.7), Inches(1.8), Inches(1.5)])

    create_styled_heading2(doc, "2.2 Software & Framework Stack")
    add_bullet_p(doc, " Next.js 14.2+ (App Router), TypeScript 5.0+, TailwindCSS (Custom Glassmorphism Tokens), Mapbox GL JS v3, Deck.gl v8, Lucide React Icons, Playwright Browser Automation Suite.", "• Frontend Environment:")
    add_bullet_p(doc, " Python 3.11+, FastAPI (Uvicorn ASGI Server), LangGraph, LangChain Core, Google Gemini 2.5 Flash / Claude / DeepSeek R1 (via NVIDIA NIM), NetworkX, Geopy, PostGIS 3.3+.", "• Backend Environment:")
    add_bullet_p(doc, " PostgreSQL 15+ dengan ekstensi spatial PostGIS 3.3+ (Supabase Managed Layer), Redis 7.0+ (Local Redis atau Upstash Serverless Redis).", "• Database & Caching:")

    # -------------------------------------------------------------
    # BAB 3: PANDUAN INSTALASI & DEPLOYMENT
    # -------------------------------------------------------------
    create_styled_heading1(doc, "BAB 3: PANDUAN INSTALASI & DEPLOYMENT LINGKUNGAN")
    
    create_styled_heading2(doc, "3.1 Kloning Repositori")
    add_code_block(doc, "git clone https://github.com/Zhav1/peta-nadi.git prehub\ncd prehub")

    create_styled_heading2(doc, "3.2 Konfigurasi Environment Variables")
    add_body_p(doc, "A. Konfigurasi Backend (`backend/.env`):")
    add_code_block(doc, 
        "APP_NAME=\"PreHub API\"\n"
        "ENVIRONMENT=\"production\"\n"
        "PORT=8000\n"
        "HOST=\"0.0.0.0\"\n"
        "GOOGLE_API_KEY=\"your-gemini-api-key\"\n"
        "TOMTOM_API_KEY=\"your-tomtom-api-key\"\n"
        "BMKG_API_URL=\"https://data.bmkg.go.id/DataMKG/TEWS/\"\n"
        "DATABASE_URL=\"postgresql://postgres:password@localhost:5432/prehub\"\n"
        "REDIS_URL=\"redis://localhost:6379/0\"\n"
        "SUPABASE_URL=\"https://your-project.supabase.co\"\n"
        "SUPABASE_SERVICE_ROLE_KEY=\"your-supabase-key\""
    )

    add_body_p(doc, "B. Konfigurasi Frontend (`frontend/.env.local`):")
    add_code_block(doc,
        "NEXT_PUBLIC_MAPBOX_TOKEN=pk.eyJ1IjoicWhhbmFraW56aGF2aSIsImEiOiJjbXI4cG8zN2wxazE5MnhweGwweHY0d2F2In0.rdp0gPLafjh-8X3IZttVog\n"
        "NEXT_PUBLIC_SUPABASE_URL=https://ulpmmacsdkohwkmyhlwj.supabase.co\n"
        "NEXT_PUBLIC_SUPABASE_ANON_KEY=your-supabase-anon-key\n"
        "NEXT_PUBLIC_API_URL=http://localhost:8000\n"
        "NEXT_PUBLIC_WS_URL=ws://localhost:8000"
    )

    create_styled_heading2(doc, "3.3 Menjalankan Backend Service (FastAPI)")
    add_code_block(doc, 
        "cd backend\n"
        "python -m venv .venv\n"
        "# Windows: .venv\\Scripts\\activate | Linux/macOS: source .venv/bin/activate\n"
        "pip install -r requirements.txt\n"
        "pytest  # Verifikasi seluruh 34 unit & integration test lulus 100%\n"
        "uvicorn app.main:app --host 0.0.0.0 --port 8000 --reload"
    )

    create_styled_heading2(doc, "3.4 Menjalankan Frontend Web Application (Next.js)")
    add_code_block(doc,
        "cd frontend\n"
        "npm install\n"
        "npm run build   # Kompilasi production build & dynamic chunks\n"
        "npm run start   # Menjalankan Next.js Production Server di port 3000"
    )

    create_styled_heading2(doc, "3.5 Otomasi Tangkapan Layar Aplikasi (Playwright)")
    add_code_block(doc, "cd frontend\nnpx playwright test e2e/capture-screenshots.spec.ts")

    # -------------------------------------------------------------
    # BAB 4: ARSITEKTUR STRUKTURAL & FORMULASI MATEMATIKA
    # -------------------------------------------------------------
    create_styled_heading1(doc, "BAB 4: ARSITEKTUR STRUKTURAL & FORMULASI MATEMATIKA")
    
    create_styled_heading2(doc, "4.1 Topologi 6 Multi-Agent Swarm")
    add_bullet_p(doc, " Melakukan deduplikasi hash, validasi telemetri BMKG, Open-Meteo, dan TomTom.", "1. Data Collection & Health Agent:")
    add_bullet_p(doc, " Mengagregasi Google News RSS dengan Source Reliability Scoring (0.0-1.0) dan pemetaan poligon bahaya.", "2. OSINT & Intelligence Agent:")
    add_bullet_p(doc, " Memproyeksikan presipitasi hujan 24-48 jam dan kemacetan segmen jalan arteri.", "3. Congestion & Weather Forecast Agent:")
    add_bullet_p(doc, " Mengoptimasi graf rute multimoda (darat, laut, udara) menggunakan algoritma NetworkX Dijkstra berpenalti zona bahaya.", "4. Logistics & Multi-Modal Routing Agent:")
    add_bullet_p(doc, " Mendeteksi anomali z-score harga pangan (cabai, beras, minyak) dan efek keterlambatan pasokan.", "5. Price & Inflation Intelligence Agent:")
    add_bullet_p(doc, " Menghasilkan sintesis penalaran mendalam (Chain-of-Thought), matriks mitigasi 3 arah, serta draf rencana aksi gabungan.", "6. Decision Support Copilot (DeepSeek R1):")

    create_styled_heading2(doc, "4.2 Formulasi Matematika Indeks Risiko Gabungan & Optimasi Rute")
    add_body_p(doc, "A. Komputasi Probabilitas Gangguan Gabungan (P_disruption):", "Formulasi 1:")
    add_code_block(doc, "P_disruption(s) = 1 - ( (1 - w_W * p_W(s)) * (1 - w_T * p_T(s)) * (1 - w_I * p_I(s)) )\n\nDi mana:\n• p_W(s) : Probabilitas risiko cuaca BMKG / Open-Meteo (bobot w_W = 0.35)\n• p_T(s) : Probabilitas kemacetan & insiden TomTom (bobot w_T = 0.40)\n• p_I(s) : Probabilitas validitas laporan OSINT berita (bobot w_I = 0.25)")

    add_body_p(doc, "B. Total Skor Risiko Operasional (R):", "Formulasi 2:")
    add_code_block(doc, "R = P_disruption(s) * ( alpha * Delta_T_delay + beta * Delta_C_fuel + gamma * V_perishability )\n\nDi mana alpha, beta, gamma adalah koefisien sensitivitas waktu, biaya bahan bakar, dan faktor risiko kebusukan muatan pangan basah.")

    add_body_p(doc, "C. Matriks Keputusan Mitigasi Tiga Arah (Tri-Option Mitigation Matrix):", "Formulasi 3:")
    add_bullet_p(doc, " Diterapkan jika R_current > R_threshold dan Biaya(Detour) < Kerugian(Spoilage/Failure).", "• REROUTE:")
    add_bullet_p(doc, " Diterapkan jika seluruh rute alternatif memiliki R_alt > R_critical (jalur terisolasi), armada ditahan di buffer depot terdekat.", "• HOLD / DELAY:")
    add_bullet_p(doc, " Diterapkan jika R_current <= R_threshold dengan rekomendasi panduan kecepatan aman (speed advisory).", "• CONTINUE:")

    # -------------------------------------------------------------
    # BAB 5: DESKRIPSI FUNGSIONAL MODUL SISTEM
    # -------------------------------------------------------------
    create_styled_heading1(doc, "BAB 5: DESKRIPSI FUNGSIONAL MODUL SISTEM")
    
    create_styled_heading2(doc, "5.1 Modul Ingesti Data Multi-Sumber & Grounding Real-Time")
    add_body_p(doc, "Modul ini bertugas menarik data primer cuaca, kemacetan, dan intelijen berita secara asynchronous setiap 1-5 menit. Data diverifikasi keabsahannya sebelum dialirkan ke Redis stream lrip:stream:osint dan disimpan ke basis data PostGIS.")

    create_styled_heading2(doc, "5.2 Modul Peta Komando 4D & Dynamic Fleet Layer")
    add_body_p(doc, "Menampilkan visualisasi spasial 60 FPS untuk seluruh armada logistik di Pulau Sumatera. Dilengkapi dengan filter modalitas interaktif (All, Land, Sea, Air), marker hub adaptif zoom level, visualisasi jalur laut Selat Malaka/Sunda, serta overlay perbandingan rute eksisting vs rute bypass mitigasi.")

    create_styled_heading2(doc, "5.3 Modul Analisis Spasial & Causal Chain Graph")
    add_body_p(doc, "Menyajikan aliran komoditas pangan makro Nusantara dengan Deck.gl Arc Layer serta analisis kausal keterlambatan distribusi terhadap lonjakan harga pangan lokal (PIHPS Grounding).")

    create_styled_heading2(doc, "5.4 Modul Multi-Agency Simulation Sandbox (What-If Advisor)")
    add_body_p(doc, "Memungkinkan pengambil kebijakan menguji skenario dampak bencana kustom (radius 5-50 km, tingkat keparahan, jenis komoditas) dan secara otomatis merumuskan Unified Action Plan lintas instansi (BAPANAS, KEMENHUB, BULOG, DISHUB/POLRI).")

    create_styled_heading2(doc, "5.5 Modul B2G Executive Cabinet Briefing Center")
    add_body_p(doc, "Menghasilkan dokumen taktis berkas kabinet berformat standar kementerian berbasis penalaran DeepSeek R1 dengan opsi Print PDF resmi, unduh JSON Telemetri, dan integrasi WhatsApp Dispatcher.")

    # -------------------------------------------------------------
    # BAB 6: PANDUAN OPERASIONAL PENGGUNA (USER MANUAL & SOP)
    # -------------------------------------------------------------
    create_styled_heading1(doc, "BAB 6: PANDUAN OPERASIONAL PENGGUNA (USER MANUAL & SOP)")
    
    sop_headers = ["No", "Tahapan Operasional", "Deskripsi Tindakan Dispatcher / Operator", "Hasil Sistem (Output)"]
    sop_rows = [
        ["1", "Akses Platform", "Buka browser ke URL web PreHub (http://localhost:3000 atau Vercel URL)", "Halaman Onboarding & Kinetic Feature Grid terbuka"],
        ["2", "Buka Command Center", "Klik tombol 'LAUNCH COMMAND CENTER 4D'", "Peta 4D & Top Nav Telemetry aktif 60 FPS"],
        ["3", "Pilih Filter Modalitas", "Pilih filter All / Land / Sea / Air pada panel kontrol peta", "Armada tersaring sesuai moda transportasi yang dipilih"],
        ["4", "Monitoring Radar Insiden", "Perhatikan daftar disrupsi aktif di panel kiri (misal: Banjir Rob Belawan)", "Skor National Logistics Health & status insiden ditampilkan"],
        ["5", "Aktivasi Swarm Reasoning", "Klik tombol '▶ Run Demo'", "6 AI Agents menjalankan simulasi penalaran 8-step berurutan"],
        ["6", "Evaluasi Evidence Chain", "Buka tab 'Evidence' di Crisis Sidebar kanan", "Verifikasi data BMKG, TomTom, dan artikel OSINT terverifikasi"],
        ["7", "Eksekusi Mitigasi", "Buka tab 'Mitigation' -> Klik 'SETUJUI & TERAPKAN RUTE ALTERNATIF'", "Rute armada diperbarui di peta & disposisi terkirim ke armada"],
        ["8", "Ekspor Laporan Pimpinan", "Buka tab 'REPORTS' -> Klik 'Print PDF Briefing' atau 'Download JSON'", "Dokumen laporan kabinet resmi siap diserahkan ke pimpinan"]
    ]
    add_styled_table(doc, sop_headers, sop_rows, col_widths=[Inches(0.4), Inches(1.5), Inches(2.6), Inches(2.0)])

    # -------------------------------------------------------------
    # BAB 7: GALERI TANGKAPAN LAYAR APLIKASI
    # -------------------------------------------------------------
    create_styled_heading1(doc, "BAB 7: GALERI TANGKAPAN LAYAR APLIKASI (VISUAL VERIFICATION)")
    add_body_p(doc, "Semua tangkapan layar di bawah ini ditangkap secara otomatis menggunakan Playwright Browser Automation Suite pada resolusi Full HD (1920x1080) dari sistem PreHub yang sedang berjalan aktif.")

    ss_dir = os.path.join(project_root, "docs", "screenshots")
    
    create_styled_heading2(doc, "7.1 Halaman Onboarding & Portal Pengenalan Sistem (Hero Section)")
    add_screenshot_figure(
        doc,
        os.path.join(ss_dir, "01_onboarding_hero.png"),
        "Gambar 7.1",
        "Tampilan Hero Section PreHub Onboarding Portal dengan visualisasi koridor logistik 3D, status rute aktif, dan navigasi cepat menuju Command Center."
    )

    create_styled_heading2(doc, "7.2 Grid Fitur Interaktif Sistem (Kinetic Feature Grid)")
    add_screenshot_figure(
        doc,
        os.path.join(ss_dir, "02_onboarding_features.png"),
        "Gambar 7.2",
        "Grid Fitur Interaktif PreHub yang menyajikan 6 pilar teknologi: Multi-Source Grounding, Multi-Agent Swarm, 4D Tactical Mapping, GPU Route Optimization, Realtime Disruption Matrix, dan B2G Cabinet Reporting."
    )

    create_styled_heading2(doc, "7.3 Pusat Komando Peta Taktis 4D & Dynamic Multimodal Fleet")
    add_screenshot_figure(
        doc,
        os.path.join(ss_dir, "03_command_center_map.png"),
        "Gambar 7.3",
        "Antarmuka Peta Komando Taktis 4D PreHub Command Center menampilkan pergerakan truk darat, kapal kargo Tol Laut via Selat Malaka, dan pesawat kargo udara dengan filter modalitas."
    )

    create_styled_heading2(doc, "7.4 Radar Insiden Logistik & Pipeline Kolaborasi Multi-Agent Swarm")
    add_screenshot_figure(
        doc,
        os.path.join(ss_dir, "04_incident_radar_pipeline.png"),
        "Gambar 7.4",
        "Radar Insiden Logistik dan Status Eksekusi 6 Multi-Agent Swarm saat memproses konsensus risiko dan penelusuran bukti (Evidence Chain)."
    )

    create_styled_heading2(doc, "7.5 Analisis Spasial Ekonomi & Causal Chain Graph (Deck.gl Layer)")
    add_screenshot_figure(
        doc,
        os.path.join(ss_dir, "05_spatial_economic_analytics.png"),
        "Gambar 7.5",
        "Modul Analisis Spasial Ekonomi Nusantara dan Pemantauan Disparitas Harga Pangan Strategis (PIHPS Grounding) dengan analisis z-score lonjakan harga."
    )

    create_styled_heading2(doc, "7.6 Multi-Agency Simulation Sandbox & What-If Advisor")
    add_screenshot_figure(
        doc,
        os.path.join(ss_dir, "06_simulation_agency_sandbox.png"),
        "Gambar 7.6",
        "Antarmuka Simulasi Kebijakan Lintas Instansi (What-If Advisor) dengan shockwave radius 5-50 km dan matriks respons bersama BAPANAS, Kemenhub, Bulog, dan Dishub/Polri."
    )

    create_styled_heading2(doc, "7.7 B2G Executive Cabinet Briefing Center")
    add_screenshot_figure(
        doc,
        os.path.join(ss_dir, "07_executive_cabinet_reports.png"),
        "Gambar 7.7",
        "Tampilan Laporan Kabinet Eksekutif (B2G Cabinet Briefing Center) siap cetak PDF resmi, ekspor JSON, dan disposisi taktis."
    )

    # -------------------------------------------------------------
    # BAB 8: STRATEGI DEPLOYMENT & ARSITEKTUR CLOUD GRATIS
    # -------------------------------------------------------------
    create_styled_heading1(doc, "BAB 8: STRATEGI DEPLOYMENT & ARSITEKTUR CLOUD GRATIS ($0 / BULAN)")
    add_body_p(doc, "Untuk mengakomodasi kebutuhan demo langsung dan evaluasi penjurian tanpa biaya operasional infrastruktur, arsitektur sistem PreHub dirancang terpisah (Decoupled Frontend-Backend Architecture):")
    
    cloud_headers = ["Layer Arsitektur", "Layanan Rekomendasi Gratis", "Karakteristik & Keunggulan", "Konfigurasi Kunci"]
    cloud_rows = [
        ["Frontend Web", "Vercel (Hobby Free)", "CI/CD otomatis dari GitHub, Global Edge CDN, HTTPS otomatis, performa WebGL optimal.", "Set NEXT_PUBLIC_API_URL ke backend"],
        ["Backend API", "Koyeb (Eco Free) / Render (Free)", "Menjalankan Python FastAPI ASGI (uvicorn). Koyeb: Tidak pernah tidur (no spin-down). Render: Pasang cron ping.", "Build: pip install -r requirements.txt\nRun: uvicorn app.main:app --port $PORT"],
        ["Database Spatial", "Supabase (Free Tier)", "500 MB PostgreSQL 15 + PostGIS 3.3, REST API otomatis, backup harian.", "Gunakan DATABASE_URL Supabase"],
        ["Cache & Streams", "Upstash Redis (Free Tier)", "10.000 request/hari gratis, kompatibel protokol Redis standar, region Singapura.", "Set REDIS_URL Upstash"],
        ["LLM AI Engine", "Google AI Studio (Gemini 2.5 Flash)", "Free Tier 15 RPM, latensi sangat cepat, penalaran cerdas.", "Set GOOGLE_API_KEY"]
    ]
    add_styled_table(doc, cloud_headers, cloud_rows, col_widths=[Inches(1.3), Inches(1.7), Inches(2.2), Inches(1.3)])

    add_callout_box(
        doc,
        "Rekomendasi Hosting Backend Terbaik: Koyeb vs Render",
        "1. Koyeb Eco Free Tier: Direkomendasikan sebagai pilihan utama karena TIDAK MASUK MODE TIDUR (no sleep/spin down), sehingga juri/penilai mendapatkan respon secepat kilat saat pertama kali membuka web.\n2. Render Free Web Service: Pilihan alternatif yang sangat stabil. Jika menggunakan Render, tambahkan URL https://your-backend.onrender.com/health ke layanan cron gratis (misal cron-job.org) setiap 10 menit untuk mencegah sleep.",
        "success"
    )

    # -------------------------------------------------------------
    # BAB 9: PENANGANAN MASALAH (TROUBLESHOOTING)
    # -------------------------------------------------------------
    create_styled_heading1(doc, "BAB 9: PENANGANAN MASALAH & PEMELIHARAAN (TROUBLESHOOTING)")
    
    trouble_headers = ["Gejala Masalah", "Kemungkinan Penyebab", "Solusi Tindakan (Actionable Fix)"]
    trouble_rows = [
        ["Peta Mapbox Blank / Gelap", "Token Mapbox belum diatur / limit token habis", "Periksa variabel NEXT_PUBLIC_MAPBOX_TOKEN di .env.local. Pastikan token valid."],
        ["Koneksi API / SSE Terputus", "Backend FastAPI mati atau port 8000 terblokir", "Jalankan uvicorn app.main:app --port 8000. Uji endpoint curl http://localhost:8000/health."],
        ["Multi-Agent Demo Gagal", "API Key LLM tidak valid atau habis limit", "Masukkan GOOGLE_API_KEY aktif. Sistem otomatis beralih ke Deterministic Fallback Agents jika API luar terputus."],
        ["Database Connection Error", "Koneksi Supabase / PostgreSQL terganggu", "Verifikasi koneksi internet dan string DATABASE_URL pada backend .env."],
        ["Performa Rendering Lambat", "Hardware Acceleration browser mati", "Aktifkan Hardware Acceleration pada pengaturan browser (Settings -> System -> Use graphics acceleration)."]
    ]
    add_styled_table(doc, trouble_headers, trouble_rows, col_widths=[Inches(1.8), Inches(1.9), Inches(2.8)])

    # -------------------------------------------------------------
    # BAB 10: KESIMPULAN & ROADMAP
    # -------------------------------------------------------------
    create_styled_heading1(doc, "BAB 10: KESIMPULAN & ROADMAP PENGEMBANGAN")
    add_body_p(doc, "Sistem PreHub membuktikan bahwa sinergi Multi-Agent AI Swarm, Multi-Source Data Grounding, dan Multi-Modal Network Routing mampu mentransformasi manajemen krisis logistik pangan dari pola reaktif-manual menjadi prediktif-preskriptif otomatis.")
    add_body_p(doc, "Dengan rantai pembuktian berbasis bukti (Evidence Chain) dan rencana aksi terpadu lintas instansi (Unified Multi-Agency Action Plan), PreHub siap diimplementasikan dan diintegrasikan bersama Badan Pangan Nasional (BAPANAS), Kementerian Perhubungan, dan Perum BULOG untuk menjaga stabilitas pasokan pangan dan memperkuat kedaulatan logistik nasional.")

    # Save
    doc.save(output_path)
    print(f"[SUCCESS] Document generated successfully at: {output_path}")

if __name__ == "__main__":
    project_root = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
    output_docx = os.path.join(project_root, "docs", "Dokumen_Pendukung_PreHub_Technical_Document.docx")
    build_technical_document(output_docx, project_root)
